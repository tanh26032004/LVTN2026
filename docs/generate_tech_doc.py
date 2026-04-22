"""
Script tạo file DOCX tài liệu kỹ thuật tự động.
Chạy: python docs/generate_tech_doc.py
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ===== STYLE =====
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(13)

# Helper
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0x0C, 0x4A, 0x6E)  # Sky-900
    return h

def add_paragraph_bold(label, value):
    p = doc.add_paragraph()
    run_b = p.add_run(label)
    run_b.bold = True
    run_b.font.name = 'Times New Roman'
    run_b.font.size = Pt(13)
    run_v = p.add_run(value)
    run_v.font.name = 'Times New Roman'
    run_v.font.size = Pt(13)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h_text in enumerate(headers):
        hdr_cells[i].text = h_text
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = str(cell_text)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
    doc.add_paragraph()

# =====================================================
# TRANG BÌA
# =====================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('TÀI LIỆU KỸ THUẬT HỆ THỐNG')
run.bold = True
run.font.size = Pt(22)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0x0C, 0x4A, 0x6E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run('Hệ thống Gợi ý Chuyên ngành Đại học\ndựa trên Phân lớp dữ liệu đa yếu tố:\nTính cách (MBTI) và Năng lực Học tập')
run2.font.size = Pt(16)
run2.font.name = 'Times New Roman'
run2.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = info.add_run('Dự án Khóa luận Tốt nghiệp — LVTN 2026')
run3.font.size = Pt(14)
run3.font.name = 'Times New Roman'
run3.font.italic = True

doc.add_page_break()

# =====================================================
# PHẦN 1: TỔNG QUAN CÔNG NGHỆ
# =====================================================
add_heading_styled('PHẦN 1: TỔNG QUAN CÔNG NGHỆ SỬ DỤNG', level=1)

add_heading_styled('1.1. Ngôn ngữ lập trình', level=2)
add_paragraph_bold('• Python 3.9+: ', 'Ngôn ngữ lập trình chính cho toàn bộ hệ thống, bao gồm xử lý dữ liệu, huấn luyện mô hình Machine Learning và phát triển giao diện Web Application.')

add_heading_styled('1.2. Framework Giao diện Web', level=2)
add_table(
    ['Thư viện', 'Phiên bản', 'Vai trò'],
    [
        ['Streamlit', '1.41.0', 'Framework chính để xây dựng giao diện Web App tương tác. Hỗ trợ render các thành phần UI (Tabs, Buttons, Charts, Data Editor) trực tiếp từ Python mà không cần viết HTML/JS riêng.'],
        ['HTML/CSS (Inline)', '—', 'Được nhúng trực tiếp qua st.markdown(unsafe_allow_html=True) để tùy chỉnh giao diện nâng cao: Hero Banner, Card Layout, Responsive Design, Glassmorphism.'],
    ]
)

add_heading_styled('1.3. Thư viện Khoa học Dữ liệu & Machine Learning', level=2)
add_table(
    ['Thư viện', 'Phiên bản', 'Vai trò'],
    [
        ['Pandas', '2.2.3', 'Xử lý, biến đổi và phân tích dữ liệu dạng bảng (DataFrame). Là xương sống của toàn bộ pipeline dữ liệu.'],
        ['NumPy', '1.26.4', 'Tính toán số học ma trận, hỗ trợ thêm nhiễu Gaussian trong quá trình Data Augmentation.'],
        ['Scikit-learn', '1.6.0', 'Thư viện Machine Learning cốt lõi, cung cấp các thuật toán phân lớp (RF, DT, SVM), tiền xử lý dữ liệu (StandardScaler, OneHotEncoder), tối ưu siêu tham số (GridSearchCV) và đánh giá mô hình.'],
        ['Joblib', '1.4.2', 'Lưu trữ (serialize) và tải lại các Pipeline mô hình đã huấn luyện dưới dạng file .joblib, cho phép triển khai trực tiếp lên Web mà không cần huấn luyện lại.'],
    ]
)

add_heading_styled('1.4. Thư viện Trực quan hóa Dữ liệu', level=2)
add_table(
    ['Thư viện', 'Phiên bản', 'Vai trò'],
    [
        ['Matplotlib', '3.10.0', 'Vẽ biểu đồ tầm quan trọng đặc trưng (Feature Importances) trong module XAI.'],
        ['Seaborn', '0.13.2', 'Bổ trợ Matplotlib với các bảng màu chuyên nghiệp và biểu đồ thống kê nâng cao.'],
    ]
)

add_heading_styled('1.5. Bảo mật & Quản trị', level=2)
doc.add_paragraph('• Xác thực Admin qua URL ẩn: Người quản trị truy cập trang đăng nhập bằng tham số URL bí mật (?admin=true), hoàn toàn vô hình với người dùng cuối.')
doc.add_paragraph('• Mã hóa Mật khẩu SHA-256: Mật khẩu Admin được lưu trữ dưới dạng mã băm một chiều (SHA-256 Hash) trong file .streamlit/secrets.toml, đảm bảo không thể bị đọc ngược ngay cả khi mã nguồn bị lộ.')
doc.add_paragraph('• Dashboard Quản trị tách biệt: Toàn bộ chức năng quản lý (CRUD câu hỏi MBTI, Upload hình ảnh, Thống kê lượt sử dụng, XAI) được gói gọn trong một module riêng biệt, không ảnh hưởng đến trải nghiệm người dùng.')

doc.add_page_break()

# =====================================================
# PHẦN 2: KIẾN TRÚC LÕI MÔ HÌNH AI
# =====================================================
add_heading_styled('PHẦN 2: KIẾN TRÚC LÕI MÔ HÌNH AI', level=1)

add_heading_styled('2.1. Hệ thống Gợi ý Lai (Hybrid Recommender System)', level=2)
doc.add_paragraph('Hệ thống sử dụng kiến trúc gợi ý lai (Hybrid Recommendation) kết hợp hai trường phái tiếp cận chính theo tỷ trọng cố định:')
add_table(
    ['Thành phần', 'Tỷ trọng', 'Kỹ thuật', 'Mô tả'],
    [
        ['Content-Based Filtering', '60%', 'Random Forest Classifier', 'Dự đoán nhóm ngành phù hợp dựa trên vector đặc trưng cá nhân của người dùng (điểm số + MBTI) thông qua mô hình phân lớp đã được huấn luyện.'],
        ['Collaborative Filtering', '40%', 'KNN (Cosine Similarity)', 'Tìm kiếm K=5 sinh viên có hồ sơ tương đồng nhất trong cơ sở dữ liệu lịch sử (8.000 hồ sơ), sau đó tổng hợp ngành học của nhóm "láng giềng gần" này theo trọng số rating.'],
    ]
)

add_heading_styled('2.2. Các thuật toán Phân lớp (Classification Models)', level=2)
doc.add_paragraph('Hệ thống huấn luyện và so sánh song song 3 thuật toán Machine Learning:')
add_table(
    ['Thuật toán', 'Vai trò', 'Siêu tham số chính', 'Accuracy'],
    [
        ['Random Forest Classifier', 'Mô hình chính (Primary)', 'n_estimators: [100, 200, 300], max_depth: [20, 30, 40, None], min_samples_split: [2, 5, 10], min_samples_leaf: [1, 2, 4] → Tối ưu bằng GridSearchCV (5-fold CV)', '~87.99%'],
        ['Decision Tree Classifier', 'Mô hình đối trọng / XAI', 'max_depth = 25, class_weight = "balanced"', '~82.36%'],
        ['SVM (RBF Kernel)', 'Mô hình đối trọng', 'kernel = "rbf", C = 0.05, probability = True, class_weight = "balanced"', '~86.99%'],
    ]
)

add_heading_styled('2.3. Pipeline tiền xử lý dữ liệu (Preprocessing Pipeline)', level=2)
doc.add_paragraph('Toàn bộ quá trình tiền xử lý được đóng gói trong sklearn.pipeline.Pipeline kết hợp sklearn.compose.ColumnTransformer, bao gồm:')
doc.add_paragraph('• StandardScaler: Chuẩn hóa (Z-score Normalization) toàn bộ 10 đặc trưng số (7 điểm gốc + 3 điểm phái sinh) về phân phối trung bình = 0, phương sai = 1. Điều này giúp SVM và KNN (Cosine Similarity) không bị lệch bởi thang đo khác nhau giữa các môn.')
doc.add_paragraph('• OneHotEncoder: Mã hóa nhị phân đặc trưng phân категории MBTI (16 nhóm tính cách) thành 16 cột binary (0/1), tránh tạo ra thứ tự giả (ordinal bias) khi đưa vào mô hình.')

doc.add_page_break()

# =====================================================
# PHẦN 3: TỐI ƯU HÓA MÔ HÌNH
# =====================================================
add_heading_styled('PHẦN 3: CHIẾN LƯỢC TỐI ƯU HÓA MÔ HÌNH', level=1)

add_heading_styled('3.1. Feature Engineering (Kỹ thuật trích chọn Đặc trưng)', level=2)
doc.add_paragraph('Ngoài 7 điểm số thô ban đầu và 1 trường MBTI, hệ thống tự động tính toán và bổ sung 3 đặc trưng phái sinh (Derived Features) quan trọng:')
add_table(
    ['Đặc trưng phái sinh', 'Công thức', 'Ý nghĩa'],
    [
        ['avg_score', 'Trung bình cộng 7 môn', 'Phản ánh năng lực học tập tổng thể của học sinh, giúp mô hình phân biệt rõ các nhóm "giỏi toàn diện" vs "giỏi chuyên môn".'],
        ['natural_science_score', 'TB(Toán, Lý, Hóa, Sinh)', 'Đánh giá xu hướng KHTN, hữu ích với các nhóm ngành Kỹ thuật, Y tế, CNTT.'],
        ['social_science_score', 'TB(Văn, Sử, Anh)', 'Đánh giá xu hướng KHXH, hữu ích với các nhóm ngành Luật, Ngoại ngữ, Báo chí.'],
    ]
)

add_heading_styled('3.2. Tối ưu Siêu tham số (Hyperparameter Tuning)', level=2)
doc.add_paragraph('Thuật toán Random Forest (mô hình chính) được tối ưu siêu tham số bằng kỹ thuật Grid Search Cross-Validation (GridSearchCV) với cấu hình:')
doc.add_paragraph('• Số fold Cross-Validation: 5 (5-fold Stratified CV)')
doc.add_paragraph('• Hàm mục tiêu tối ưu: F1-Score Macro (cân bằng hiệu suất giữa tất cả 12 lớp ngành)')
doc.add_paragraph('• Không gian tìm kiếm: 3 × 4 × 3 × 3 = 108 tổ hợp siêu tham số')
doc.add_paragraph('• Chế độ song song: n_jobs = -1 (sử dụng toàn bộ lõi CPU)')

add_heading_styled('3.3. Chống mất cân bằng lớp (Class Imbalance Handling)', level=2)
doc.add_paragraph('Tất cả 3 mô hình đều được cấu hình class_weight="balanced", cho phép Scikit-learn tự động điều chỉnh trọng số của từng lớp theo tỷ lệ nghịch với tần suất xuất hiện. Điều này đảm bảo các nhóm ngành có ít mẫu (ví dụ: Nghệ thuật, Nông nghiệp) không bị "lấn át" bởi các nhóm phổ biến hơn (CNTT, Kinh tế).')

doc.add_page_break()

# =====================================================
# PHẦN 4: DỮ LIỆU
# =====================================================
add_heading_styled('PHẦN 4: DỮ LIỆU HUẤN LUYỆN', level=1)

add_heading_styled('4.1. Nguồn dữ liệu gốc (Raw Data Sources)', level=2)
add_table(
    ['STT', 'Nguồn dữ liệu', 'Mô tả', 'Liên kết'],
    [
        ['1', 'VN Student Performance Dataset', 'Dữ liệu điểm thi THPT Quốc gia Việt Nam', 'Kaggle (kaggle.com/datasets/hongngctin/vn-student-performance-dataset)'],
        ['2', 'Career Path Recommendation (Indonesia)', 'Dữ liệu phân tích tính cách MBTI và điểm số học sinh Indonesia', 'Mendeley Data (data.mendeley.com/datasets/yzbpwk2wnf/1)'],
        ['3', 'Bài báo khoa học Data in Brief (Elsevier)', 'Tài liệu tham khảo và chứng minh phương pháp luận', 'DOI: 10.1016/j.dib.2025.111438'],
    ]
)

add_heading_styled('4.2. Quy trình xử lý và Tăng cường Dữ liệu (Data Pipeline)', level=2)
doc.add_paragraph('Dữ liệu thô từ 3 nguồn trên được xử lý qua quy trình gồm 3 bước:')
doc.add_paragraph('Bước 1 — Ghép nối dữ liệu (merge_raw_data.py): Kết hợp các nguồn dữ liệu khác nhau bằng thuật toán KNN matching, đồng bộ hóa thang điểm và tạo ra bộ dữ liệu gốc thống nhất.')
doc.add_paragraph('Bước 2 — Tăng cường dữ liệu (data_augmentation.py): Sử dụng kỹ thuật Data Augmentation với nhiễu Gaussian (σ = 0.6) và tịnh tiến phổ điểm (Mean Shift) theo điểm trung bình THPT Quốc gia 2023 do Bộ GD-ĐT công bố, nhân bản dữ liệu lên gấp 5 lần.')
doc.add_paragraph('Bước 3 — Huấn luyện mô hình (train_model.py): Chia tập dữ liệu theo tỷ lệ 80:20 (Train:Test) với Stratified Split để đảm bảo phân phối lớp đầu ra cân đối.')

add_heading_styled('4.3. Thông số bộ dữ liệu cuối cùng', level=2)
add_table(
    ['Đặc điểm', 'Giá trị'],
    [
        ['Tổng số mẫu (Samples)', '~8.000 hồ sơ học sinh'],
        ['Số lớp đầu ra (Target Classes)', '12 nhóm chuyên ngành đại học'],
        ['Số đặc trưng gốc (Raw Features)', '7 điểm số + 1 MBTI type = 8'],
        ['Số đặc trưng sau Feature Engineering', '10 số + 16 binary (OneHot MBTI) = 26 chiều'],
        ['Tỷ lệ chia Train/Test', '80% / 20% (Stratified)'],
        ['Phương pháp Augmentation', 'Gaussian Noise (σ=0.6) + Mean Shift theo phổ điểm BGD 2023'],
    ]
)

add_heading_styled('4.4. Danh sách 12 Nhóm ngành Phân lớp', level=2)
majors = [
    'CNTT & Kỹ thuật Máy tính', 'Kinh tế & Quản lý', 'Y tế & Sức khỏe',
    'Sư phạm & Giáo dục', 'Luật & Chính trị', 'Ngoại ngữ & Ngôn ngữ',
    'Nghệ thuật & Thiết kế', 'Kỹ thuật & Công nghệ', 'Khoa học Tự nhiên',
    'Khoa học Xã hội & Nhân văn', 'Nông Lâm Ngư nghiệp', 'Báo chí & Truyền thông'
]
for i, m in enumerate(majors, 1):
    doc.add_paragraph(f'{i}. {m}')

doc.add_page_break()

# =====================================================
# PHẦN 5: KIẾN TRÚC MÃ NGUỒN
# =====================================================
add_heading_styled('PHẦN 5: KIẾN TRÚC MÃ NGUỒN', level=1)

doc.add_paragraph('Hệ thống được thiết kế theo kiến trúc Module MVC (Model-View-Controller) phân tách rõ ràng giữa tầng dữ liệu, tầng xử lý nghiệp vụ và tầng giao diện:')

add_table(
    ['Thư mục / File', 'Chức năng'],
    [
        ['app.py', 'Điểm khởi chạy chính, phân luồng Admin/User qua query params, render Hero Banner và Tab navigation.'],
        ['views/components.py', 'Module CSS/HTML chung: Custom Styles, Hero Banner, Pill Tabs, Responsive Design.'],
        ['views/tab1_survey.py', 'Giao diện Khảo sát & Dự đoán Ngành học: nhập MBTI, chọn khối thi, nhập điểm → gọi Hybrid Recommender.'],
        ['views/tab2_mbti.py', 'Giao diện Trắc nghiệm MBTI 16 câu hỏi tương tác với kết quả thời gian thực.'],
        ['views/tab3_xai.py', 'Module Giải thích AI (XAI): biểu đồ Feature Importances từ Random Forest.'],
        ['views/admin_dashboard.py', 'Dashboard quản trị ẩn: Chọn thuật toán, CRUD câu hỏi MBTI, Upload ảnh, Thống kê lượt sử dụng.'],
        ['scripts/train_model.py', 'Pipeline huấn luyện 3 mô hình ML với Feature Engineering, GridSearchCV, xuất .joblib.'],
        ['scripts/hybrid_recommender.py', 'Module thuật toán Gợi ý Lai (Content-Based 60% + Collaborative Filtering 40%).'],
        ['scripts/data_augmentation.py', 'Script tăng cường dữ liệu: Gaussian Noise + Mean Shift theo phổ điểm BGD.'],
        ['data/student_data.csv', 'Bộ dữ liệu huấn luyện chính (8.000 hồ sơ).'],
        ['data/questions/mbti_questions.json', 'Cơ sở dữ liệu câu hỏi trắc nghiệm MBTI (quản lý qua Admin Dashboard).'],
        ['model/*.joblib', 'Các file Pipeline mô hình đã huấn luyện: RF, DT, SVM, Preprocessor, Target Encoder.'],
        ['.streamlit/secrets.toml', 'Lưu trữ mã băm SHA-256 mật khẩu Admin (bảo mật).'],
    ]
)

# =====================================================
# LƯU FILE
# =====================================================
output_path = os.path.join(os.path.dirname(__file__), 'tai_lieu_ky_thuat_he_thong.docx')
doc.save(output_path)
print(f"✅ Đã tạo thành công file tài liệu: {output_path}")
